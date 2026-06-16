import hashlib
import random
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
import customtkinter as ctk

try:
    import openpyxl
    from docx import Document
    HAS_OFFICE_LIBS = True
except ImportError:
    HAS_OFFICE_LIBS = False
print("HAS_OFFICE_LIBS =", HAS_OFFICE_LIBS)
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class ThuatToanRSA:
    @staticmethod
    def uoc_chung_lon_nhat(a, b):
        while b:
            a, b = b, a % b
        return a
    
    @staticmethod
    def kiem_tra_so_nguyen_to(n, k=5):
        if n < 2: return False
        if n in (2, 3): return True
        if n % 2 == 0: return False
        r, s = n - 1, 0
        while r % 2 == 0:
            r //= 2
            s += 1
        for _ in range(k):
            a = random.randint(2, n - 2)
            x = pow(a, r, n)
            if x == 1 or x == n - 1: continue
            for _ in range(s - 1):
                x = pow(x, 2, n)
                if x == n - 1: break
            else: return False
        return True

    @staticmethod
    def sinh_so_nguyen_to(min_val, max_val):
        while True:
            n = random.randint(min_val, max_val)
            if n % 2 != 0 and ThuatToanRSA.kiem_tra_so_nguyen_to(n): 
                return n

    @staticmethod
    def tim_khoa_cong_khai_e(phi, so_luong=1):
        goi_y = []
        min_e = 1 << 64
        max_e = 1 << 128
        vong_lap = 0
        while len(goi_y) < so_luong and vong_lap < 2000:
            vong_lap += 1
            test_e = random.getrandbits(128)
            if min_e < test_e < phi and test_e % 2 != 0:
                if ThuatToanRSA.uoc_chung_lon_nhat(test_e, phi) == 1:
                    goi_y.append(test_e)
                    break
        if not goi_y:
            goi_y.append(65537) 
        return goi_y

    @staticmethod
    def tim_nghich_dao_modulo(e, phi):
        m0, y, x = phi, 0, 1
        while e > 1:
            q = e // phi
            t = phi
            phi = e % phi
            e = t
            t = y
            y = x - q * y
            x = t
        if x < 0: x += m0
        return x

    @staticmethod
    def doc_noi_dung_file(filepath):
        ext = filepath.split(".")[-1].lower()
        if ext == "pdf":
            import PyPDF2
            text = ""
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        ext = filepath.split(".")[-1].lower()
        if ext == "docx" and HAS_OFFICE_LIBS:
            try:
                doc = Document(filepath)
                return "\n".join([p.text for p in doc.paragraphs])
            except: pass
        elif ext == "xlsx" and HAS_OFFICE_LIBS:
            try:
                wb = openpyxl.load_workbook(filepath, data_only=True)
                gop_text = []
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        dong_str = [str(c) for c in row if c is not None]
                        if dong_str: gop_text.append(" ".join(dong_str))
                return "\n".join(gop_text)
            except: pass
        try:
            with open(filepath, "r", encoding="utf-8") as f: return f.read()
        except:
            with open(filepath, "rb") as f: return f.read().decode("utf-8", errors="ignore")

class UngDungChuKySoRSA:
    def __init__(self, root):
        self.root = root
        self.root.title(" RSA ")
        self.root.geometry("1300x900")
        self.root.minsize(1200, 850)
        self.bg_main = ("#F3F4F6", "#0B0F19")
        self.bg_sidebar = ("#E5E7EB", "#111827")
        self.bg_card = ("#FFFFFF", "#1F2937")
        self.border_card = ("#D1D5DB", "#374151")
        self.text_title = ("#111827", "#FFFFFF")
        self.accent_blue = "#3B82F6"
        self.accent_green = "#10B981"
        self.text_muted = "#9CA3AF"
        self.root.configure(fg_color=self.bg_main)
        self.n, self.e, self.d = 0, 0, 0
        self.dem_ls = 0
        self.data_ls = {}
        self.root.columnconfigure(1, weight=1)
        self.root.rowconfigure(0, weight=1)
        self.tao_thanh_menu()
        self.tao_giao_dien_chinh()

    def tao_thanh_menu(self):
        self.sidebar = ctk.CTkFrame(self.root, width=260, corner_radius=0, fg_color=self.bg_sidebar, border_width=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_propagate(False)
        ctk.CTkLabel(self.sidebar, text="🔑 KHÓA RSA", font=ctk.CTkFont(family="Inter", size=22, weight="bold"), 
                     text_color=self.text_title).grid(row=0, column=0, padx=25, pady=(40, 30), sticky="w")
        self.menus = {}
        cau_hinh_menu = [
            (1, "🔑  Quản Lý Cặp Khóa"),
            (2, "✍️  Tạo Chữ Ký Số"),
            (3, "🛡️  Thẩm Định Chữ Ký")
        ]
        for index, text in cau_hinh_menu:
            btn = ctk.CTkButton(
                self.sidebar, text=text, anchor="w", height=45, corner_radius=10,
                font=ctk.CTkFont(family="Inter", size=13, weight="bold" if index == 1 else "normal"),
                fg_color=self.accent_blue if index == 1 else "transparent",
                text_color=self.text_title,
                hover_color=("#CDD5E0", "#2D3748"),
                command=lambda idx=index: self.chuyen_tab(idx)
            )
            btn.grid(row=index, column=0, padx=15, pady=8, sticky="ew")
            self.menus[index] = btn

    def tao_giao_dien_chinh(self):
        self.container = ctk.CTkFrame(self.root, corner_radius=0, fg_color="transparent")
        self.container.grid(row=0, column=1, sticky="nsew", padx=30, pady=30)
        self.container.columnconfigure(0, weight=1)
        self.container.rowconfigure(1, weight=1)
        f_top_bar = ctk.CTkFrame(self.container, fg_color="transparent")
        f_top_bar.grid(row=0, column=0, sticky="ew", pady=(0, 25))
        f_top_bar.columnconfigure(0, weight=1)
        self.title_page = ctk.CTkLabel(f_top_bar, text="QUẢN LÝ HỆ THỐNG CẶP KHÓA RSA", font=ctk.CTkFont(family="Inter", size=24, weight="bold"), text_color=self.text_title)
        self.title_page.grid(row=0, column=0, sticky="w")
        self.switch_var = ctk.StringVar(value="dark")
        self.theme_switch = ctk.CTkSwitch(
            f_top_bar, text="Chế độ Tối", font=ctk.CTkFont(family="Inter", size=12),
            command=self.doi_giao_dien_sang_toi, variable=self.switch_var,
            onvalue="dark", offvalue="light", text_color=self.text_title,
            progress_color=self.accent_blue
        )
        self.theme_switch.grid(row=0, column=1, sticky="e", padx=5)
        self.views = {
            1: self.tao_tab_quan_ly_khoa(),
            2: self.tao_tab_ky_so(),
            3: self.tao_tab_xac_minh()
        }
        self.views[1].grid(row=1, column=0, sticky="nsew")
        self.view_hien_tai = self.views[1]

    def doi_giao_dien_sang_toi(self):
        if self.switch_var.get() == "dark":
            ctk.set_appearance_mode("Dark")
            self.theme_switch.configure(text="Chế độ Tối")
        else:
            ctk.set_appearance_mode("Light")
            self.theme_switch.configure(text="Chế độ Sáng")

    def chuyen_tab(self, id_man_hinh):
        self.view_hien_tai.grid_forget()
        tieu_de_trang = {
            1: "QUẢN LÝ HỆ THỐNG CẶP KHÓA RSA",
            2: "TẠO CHỮ KÝ SỐ ",
            3: "THẨM ĐỊNH CHỮ KÝ SỐ"
        }
        self.title_page.configure(text=tieu_de_trang[id_man_hinh])
        for idx, btn in self.menus.items():
            if idx == id_man_hinh:
                btn.configure(fg_color=self.accent_blue, font=ctk.CTkFont(family="Inter", size=13, weight="bold"))
            else:
                btn.configure(fg_color="transparent", font=ctk.CTkFont(family="Inter", size=13, weight="normal"))
        self.views[id_man_hinh].grid(row=1, column=0, sticky="nsew")
        self.view_hien_tai = self.views[id_man_hinh]

    def tao_tab_quan_ly_khoa(self):
        f_main = ctk.CTkFrame(self.container, fg_color="transparent")
        f_main.columnconfigure(0, weight=1)
        f_main.rowconfigure(1, weight=1)
        c_input = ctk.CTkFrame(f_main, corner_radius=15, fg_color=self.bg_card, border_width=1, border_color=self.border_card)
        c_input.grid(row=0, column=0, sticky="ew", pady=(0, 20), padx=5)
        c_input.columnconfigure(0, weight=1)
        ctk.CTkLabel(c_input, text="⚙️ THAM SỐ ĐẦU VÀO", font=ctk.CTkFont(size=14, weight="bold"), text_color=self.accent_blue).grid(row=0, column=0, columnspan=2, sticky="w", padx=25, pady=(20, 15))
        f_left = ctk.CTkFrame(c_input, fg_color="transparent")
        f_left.grid(row=1, column=0, sticky="ew", padx=25, pady=10)
        f_left.columnconfigure(1, weight=1)
        ctk.CTkLabel(f_left, text="Số nguyên tố P:", font=ctk.CTkFont(size=12), text_color=self.text_title).grid(row=0, column=0, sticky="w", pady=8)
        self.txt_p = ctk.CTkEntry(f_left, height=35, placeholder_text="Nhập P...", border_color=self.border_card, fg_color=self.bg_main, text_color=self.text_title)
        self.txt_p.grid(row=0, column=1, padx=(15, 0), pady=8, sticky="ew")
        ctk.CTkLabel(f_left, text="Số nguyên tố Q:", font=ctk.CTkFont(size=12), text_color=self.text_title).grid(row=1, column=0, sticky="w", pady=8)
        self.txt_q = ctk.CTkEntry(f_left, height=35, placeholder_text="Nhập Q...", border_color=self.border_card, fg_color=self.bg_main, text_color=self.text_title)
        self.txt_q.grid(row=1, column=1, padx=(15, 0), pady=8, sticky="ew")
        ctk.CTkLabel(f_left, text="Số nguyên tố E:", font=ctk.CTkFont(size=12), text_color=self.text_title).grid(row=2, column=0, sticky="w", pady=8)
        self.txt_e = ctk.CTkEntry(f_left, height=35, placeholder_text="Nhập E...", border_color=self.border_card, fg_color=self.bg_main, text_color=self.text_title)
        self.txt_e.grid(row=2, column=1, padx=(15, 0), pady=8, sticky="ew")
        f_right = ctk.CTkFrame(c_input, fg_color="transparent")
        f_right.grid(row=1, column=1, sticky="n", padx=25, pady=10)
        ctk.CTkButton(f_right, text="🎲 Sinh Ngẫu Nhiên P, Q", width=220, height=40, fg_color=("#9CA3AF", "#4B5563"), command=self.sinh_ngau_nhien_p_q).pack(pady=5)
        ctk.CTkButton(f_right, text="⚡ Tìm Khóa E Khả Dụng", width=220, height=40, fg_color="transparent", border_width=1, border_color=self.border_card, text_color=self.text_title, command=self.tim_e_hop_le).pack(pady=5)
        ctk.CTkButton(f_right, text="🔒 TẠO HỆ THỐNG CẶP KHÓA", width=220, height=40, font=ctk.CTkFont(weight="bold"), fg_color=self.accent_green, text_color="white", command=self.tao_cap_khoa_rsa).pack(pady=5)
        c_output = ctk.CTkFrame(f_main, corner_radius=15, fg_color=self.bg_card, border_width=1, border_color=self.border_card)
        c_output.grid(row=1, column=0, sticky="nsew", pady=2)
        c_output.columnconfigure(0, weight=1)
        ctk.CTkLabel(c_output, text="📋 THÔNG TIN CẶP KHÓA", font=ctk.CTkFont(size=14, weight="bold"), text_color=self.accent_green).grid(row=0, column=0, sticky="w", padx=25, pady=(20, 5))
        ctk.CTkLabel(c_output, text="Modulo (N):", font=ctk.CTkFont(weight="bold")).grid(row=1, column=0, sticky="w", padx=25)
        self.txt_display_n = ctk.CTkTextbox(c_output, height=60, font=("Consolas", 12), fg_color=self.bg_main, border_width=1, corner_radius=8)
        self.txt_display_n.grid(row=2, column=0, sticky="ew", padx=25, pady=5)
        ctk.CTkLabel(c_output, text="Cặp khóa công khai (N,E):", font=ctk.CTkFont(weight="bold"), text_color=("#DC2626", "#FCA5A5")).grid(row=3, column=0, sticky="w", padx=25)
        self.txt_display_e = ctk.CTkTextbox(c_output, height=60, font=("Consolas", 12), fg_color=self.bg_main, border_width=1, corner_radius=8)
        self.txt_display_e.grid(row=4, column=0, sticky="ew", padx=25, pady=5)
        ctk.CTkLabel(c_output, text="Cặp khóa Bí Mật (N,D):", font=ctk.CTkFont(weight="bold"), text_color=("#DC2626", "#FCA5A5")).grid(row=5, column=0, sticky="w", padx=25)
        self.txt_display_d = ctk.CTkTextbox(c_output, height=60, font=("Consolas", 12), fg_color=self.bg_main, border_width=1, corner_radius=8)
        self.txt_display_d.grid(row=6, column=0, sticky="ew", padx=25, pady=5)
        f_file_actions = ctk.CTkFrame(c_output, fg_color="transparent")
        f_file_actions.grid(row=7, column=0, sticky="w", pady=20, padx=25)
        ctk.CTkButton(f_file_actions, text="📁 Nhập Khóa", width=150, height=35, fg_color=("#6B7280", "#374151"), command=self.mo_cap_khoa).pack(side="left", padx=(0, 10))
        ctk.CTkButton(f_file_actions, text="💾 Xuất Khóa", width=150, height=35, fg_color=("#6B7280", "#374151"), command=self.luu_cap_khoa).pack(side="left")
        return f_main

    def tao_tab_ky_so(self):
        f_main = ctk.CTkFrame(self.container, fg_color="transparent")
        f_main.columnconfigure(0, weight=1)
        f_main.rowconfigure(2, weight=1)
        c_doc = ctk.CTkFrame(f_main, corner_radius=15, fg_color=self.bg_card, border_width=1, border_color=self.border_card)
        c_doc.grid(row=0, column=0, sticky="ew", pady=(0, 15))
        ctk.CTkLabel(c_doc, text="📄 NỘI DUNG VĂN BẢN CẦN KÝ SỐ ", font=ctk.CTkFont(size=14, weight="bold"), text_color=self.accent_blue).pack(anchor="w", padx=25, pady=(20, 8))
        self.txt_source_text = ctk.CTkTextbox(c_doc, height=130, font=("Segoe UI", 12), border_width=1, border_color=self.border_card, fg_color=self.bg_main, corner_radius=8, text_color=self.text_title)
        self.txt_source_text.pack(fill="x", pady=5, padx=25)
        self.txt_source_text.bind("<Control-Return>", lambda event: self.ky_so_van_ban())
        f_doc_btns = ctk.CTkFrame(c_doc, fg_color="transparent")
        f_doc_btns.pack(fill="x", pady=(10, 20), padx=25)
        ctk.CTkButton(f_doc_btns, text="📁 Tải Tài Liệu Lên ", height=36, fg_color=("#6B7280", "#4B5563"), text_color="white", hover_color=("#4B5563", "#374151"), command=self.tai_van_ban_can_ky).pack(side="left", padx=(0, 12))
        ctk.CTkButton(f_doc_btns, text="💾 Lưu Văn Bản Gốc", height=36, fg_color="transparent", border_width=1, border_color=self.border_card, text_color=self.text_title, hover_color=("#E5E7EB", "#2D3748"), command=self.luu_van_ban_goc).pack(side="left")
        c_sig = ctk.CTkFrame(f_main, corner_radius=15, fg_color=self.bg_card, border_width=1, border_color=self.border_card)
        c_sig.grid(row=1, column=0, sticky="ew", pady=(0, 15))
        ctk.CTkLabel(c_sig, text="🔑 CHỮ KÝ SỐ ", font=ctk.CTkFont(size=14, weight="bold"), text_color="#F87171").pack(anchor="w", padx=25, pady=(20, 8))
        self.txt_output_sig = ctk.CTkTextbox(c_sig, height=75, font=("Consolas", 12), fg_color=self.bg_main, border_color=self.border_card, border_width=1, corner_radius=8, text_color=self.text_title)
        self.txt_output_sig.pack(fill="x", pady=5, padx=25)
        f_sig_btns = ctk.CTkFrame(c_sig, fg_color="transparent")
        f_sig_btns.pack(fill="x", pady=(10, 20), padx=25)
        ctk.CTkButton(f_sig_btns, text="💾 Lưu Mã Chữ Ký (*.txt)", height=36, fg_color=("#6B7280", "#374151"), text_color="white", hover_color=("#4B5563", "#1F2937"), command=self.luu_chu_ky).pack(side="left")
        ctk.CTkButton(f_sig_btns, text="✍️ THỰC HIỆN KÝ SỐ ĐIỆN TỬ", height=38, font=ctk.CTkFont(weight="bold"), fg_color=self.accent_blue, text_color="white", hover_color="#2563EB", command=self.ky_so_van_ban).pack(side="right")
        c_history = ctk.CTkFrame(f_main, corner_radius=15, fg_color=self.bg_card, border_width=1, border_color=self.border_card)
        c_history.grid(row=2, column=0, sticky="nsew", pady=2)
        ctk.CTkLabel(c_history, text="📋 LỊCH SỬ PHIÊN KÝ SỐ (Nhấp chọn dòng để khôi phục nhanh dữ liệu)", font=ctk.CTkFont(size=12, weight="bold"), text_color=self.text_muted).pack(anchor="w", padx=25, pady=(15, 5))
        f_header = ctk.CTkFrame(c_history, fg_color=("#E5E7EB", "#2D3748"), height=30, corner_radius=5)
        f_header.pack(fill="x", padx=25, pady=2)
        f_header.grid_propagate(False)
        f_header.columnconfigure(2, weight=1)
        ctk.CTkLabel(f_header, text="STT", font=ctk.CTkFont(size=11, weight="bold"), width=50).grid(row=0, column=0, padx=5, sticky="w")
        ctk.CTkLabel(f_header, text="Thời Gian", font=ctk.CTkFont(size=11, weight="bold"), width=100).grid(row=0, column=1, padx=5, sticky="w")
        ctk.CTkLabel(f_header, text="Đoạn Trích Nội Dung Gốc", font=ctk.CTkFont(size=11, weight="bold"), anchor="w").grid(row=0, column=2, padx=5, sticky="ew")
        ctk.CTkLabel(f_header, text="Chữ Ký Số", font=ctk.CTkFont(size=11, weight="bold"), width=250, anchor="w").grid(row=0, column=3, padx=5, sticky="w")
        self.scroll_history = ctk.CTkScrollableFrame(c_history, fg_color="transparent")
        self.scroll_history.pack(fill="both", expand=True, padx=25, pady=(0, 15))
        return f_main

    def tao_tab_xac_minh(self):
        f_main = ctk.CTkFrame(self.container, fg_color="transparent")
        f_main.columnconfigure(0, weight=1)
        c_pub = ctk.CTkFrame(f_main, corner_radius=15, fg_color=self.bg_card, border_width=1, border_color=self.border_card)
        c_pub.grid(row=0, column=0, sticky="ew", pady=(0, 15))
        c_pub.columnconfigure(1, weight=1) 
        ctk.CTkLabel(c_pub, text="🔑 CẤU HÌNH KHÓA ĐỐI CHỨNG ", font=ctk.CTkFont(size=14, weight="bold"), text_color=self.accent_blue).grid(row=0, column=0, columnspan=2, sticky="w", padx=25, pady=(20, 12))
        ctk.CTkLabel(c_pub, text="Modulo N:", font=ctk.CTkFont(size=12), text_color=self.text_title).grid(row=1, column=0, sticky="w", padx=(25, 10), pady=10)
        self.txt_t3_n = ctk.CTkEntry(c_pub, height=35, placeholder_text="Nhập Modulo N...", border_color=self.border_card, fg_color=self.bg_main, text_color=self.text_title)
        self.txt_t3_n.grid(row=1, column=1, padx=(0, 25), pady=10, sticky="ew")
        ctk.CTkLabel(c_pub, text="Chỉ số E:", font=ctk.CTkFont(size=12), text_color=self.text_title).grid(row=2, column=0, sticky="w", padx=(25, 10), pady=(0, 20))
        self.txt_t3_e = ctk.CTkEntry(c_pub, height=35, placeholder_text="Nhập E...", border_color=self.border_card, fg_color=self.bg_main, text_color=self.text_title)
        self.txt_t3_e.grid(row=2, column=1, padx=(0, 25), pady=(0, 20), sticky="ew")
        c_check = ctk.CTkFrame(f_main, corner_radius=15, fg_color=self.bg_card, border_width=1, border_color=self.border_card)
        c_check.grid(row=1, column=0, sticky="ew", pady=(0, 15))
        ctk.CTkLabel(c_check, text="📄 VĂN BẢN VÀ CHỮ KÝ CẦN THẨM ĐỊNH", font=ctk.CTkFont(size=14, weight="bold"), text_color="#61AFEF").pack(anchor="w", padx=25, pady=(20, 8))
        self.txt_t3_text = ctk.CTkTextbox(c_check, height=90, border_width=1, border_color=self.border_card, fg_color=self.bg_main, corner_radius=8, text_color=self.text_title)
        self.txt_t3_text.pack(fill="x", pady=5, padx=25)
        f_f1 = ctk.CTkFrame(c_check, fg_color="transparent")
        f_f1.pack(fill="x", padx=25, pady=(5, 15))
        ctk.CTkButton(f_f1, text="📁 Tải Văn Bản Thẩm Định", height=28, fg_color=("#6B7280", "#4B5563"), text_color="white", hover_color=("#4B5563", "#374151"), command=self.tai_van_ban_xac_minh).pack(side="left")
        self.txt_t3_sig = ctk.CTkTextbox(c_check, height=60, border_width=1, border_color=self.border_card, fg_color=self.bg_main, corner_radius=8, text_color=self.text_title)
        self.txt_t3_sig.pack(fill="x", pady=5, padx=25)
        self.txt_t3_n.bind("<Return>", lambda event: self.xac_minh_chu_ky())
        self.txt_t3_e.bind("<Return>", lambda event: self.xac_minh_chu_ky())
        self.txt_t3_text.bind("<Control-Return>", lambda event: self.xac_minh_chu_ky())
        self.txt_t3_sig.bind("<Control-Return>", lambda event: self.xac_minh_chu_ky())
        f_f2 = ctk.CTkFrame(c_check, fg_color="transparent")
        f_f2.pack(fill="x", padx=25, pady=(5, 20))
        ctk.CTkButton(f_f2, text="📁 Tải Chữ Ký (.txt)", height=28, fg_color=("#6B7280", "#4B5563"), text_color="white", hover_color=("#4B5563", "#374151"), command=self.tai_chu_ky_xac_minh).pack(side="left")
        ctk.CTkButton(f_main, text="🛡️ BẮT ĐẦU THẨM ĐỊNH TÍNH TOÀN VẸN ", 
                      font=ctk.CTkFont(size=14, weight="bold"), height=48, corner_radius=12,
                      fg_color=self.accent_blue, text_color="white", hover_color="#2563EB", command=self.xac_minh_chu_ky).grid(row=2, column=0, sticky="ew", pady=(0, 15))
        c_log = ctk.CTkFrame(f_main, corner_radius=15, fg_color=self.bg_card, border_width=1, border_color=self.border_card)
        c_log.grid(row=3, column=0, sticky="nsew", pady=2)
        ctk.CTkLabel(c_log, text="⚙️ NHẬT KÝ CHI TIẾT BƯỚC TÍNH TOÁN GIẢI THUẬT RSA", font=ctk.CTkFont(size=12, weight="bold"), text_color=self.text_muted).pack(anchor="w", padx=25, pady=(15, 6))
        self.txt_log = ctk.CTkTextbox(c_log, font=("Consolas", 12), fg_color=self.bg_main, border_color=self.border_card, border_width=1, corner_radius=8, text_color=self.text_title)
        self.txt_log.pack(fill="both", expand=True, pady=(0, 20), padx=25)
        return f_main
    
    def sinh_ngau_nhien_p_q(self):
        min_val = 1 << 130 
        max_val = (1 << 132) - 1 
        p = ThuatToanRSA.sinh_so_nguyen_to(min_val, max_val)
        while True:
            q = ThuatToanRSA.sinh_so_nguyen_to(min_val, max_val)
            if p != q: break
        self.txt_p.delete(0, tk.END); self.txt_p.insert(0, str(p))
        self.txt_q.delete(0, tk.END); self.txt_q.insert(0, str(q))

    def tim_e_hop_le(self):
        str_p, str_q = self.txt_p.get().strip(), self.txt_q.get().strip()
        if not str_p or not str_q:
            messagebox.showwarning("Thiếu tham số", "Vui lòng chuẩn bị giá trị P, Q trước!")
            return
        try:
            phi = (int(str_p) - 1) * (int(str_q) - 1)
            list_e = ThuatToanRSA.tim_khoa_cong_khai_e(phi, 1)
            if list_e:
                self.txt_e.delete(0, tk.END); self.txt_e.insert(0, str(list_e[0]))
        except ValueError:
            messagebox.showerror("Lỗi", "Định dạng P, Q không phải số hợp lệ.")

    def tao_cap_khoa_rsa(self):
        str_p = self.txt_p.get().strip()
        str_q = self.txt_q.get().strip()
        str_e = self.txt_e.get().strip()
        if not str_p:
            messagebox.showwarning("Thiếu dữ liệu", "Vui lòng nhập số nguyên tố P!")
            self.txt_p.focus_set()
            return
        if not str_q:
            messagebox.showwarning("Thiếu dữ liệu", "Vui lòng nhập số nguyên tố Q!")
            self.txt_q.focus_set()
            return
        if not str_e:
            messagebox.showwarning("Thiếu dữ liệu", "Vui lòng nhập khóa công khai E!")
            self.txt_e.focus_set()
            return
        try:
            p = int(str_p)
        except ValueError:
            messagebox.showerror("Lỗi dữ liệu", "Giá trị P phải là số nguyên!")
            self.txt_p.focus_set()
            return
        try:
            q = int(str_q)
        except ValueError:
            messagebox.showerror("Lỗi dữ liệu", "Giá trị Q phải là số nguyên!")
            self.txt_q.focus_set()
            return
        try:
            e_val = int(str_e)
        except ValueError:
            messagebox.showerror("Lỗi dữ liệu", "Giá trị E phải là số nguyên!")
            self.txt_e.focus_set()
            return
        if not ThuatToanRSA.kiem_tra_so_nguyen_to(p):
            messagebox.showerror("Lỗi", "Giá trị P không phải số nguyên tố!")
            self.txt_p.focus_set()
            return
        if not ThuatToanRSA.kiem_tra_so_nguyen_to(q):
            messagebox.showerror("Lỗi", "Giá trị Q không phải số nguyên tố!")
            self.txt_q.focus_set()
            return
        if p == q:
            messagebox.showerror("Lỗi", "P và Q phải khác nhau!")
            return
        phi = (p - 1) * (q - 1)
        if e_val <= 1 or e_val >= phi:
            messagebox.showerror(
                "Lỗi khóa E",
                f"Giá trị E phải nằm trong khoảng (1 ; {phi})"
            )
            self.txt_e.focus_set()
            return
        if ThuatToanRSA.uoc_chung_lon_nhat(e_val, phi) != 1:
            goi_y = ", ".join(
                str(x) for x in ThuatToanRSA.tim_khoa_cong_khai_e(phi, 3)
            )
            messagebox.showerror(
                "Lỗi khóa E",
                f"Khóa E không nguyên tố cùng nhau với Phi(N)!\n\n"
                f"Gợi ý E hợp lệ:\n{goi_y}"
            )
            self.txt_e.focus_set()
            return
        self.n = p * q
        self.e = e_val
        self.d = ThuatToanRSA.tim_nghich_dao_modulo(self.e, phi)
        self.txt_display_n.delete("1.0", tk.END)
        self.txt_display_n.insert(tk.END, str(self.n))

        self.txt_display_e.delete("1.0", tk.END)
        self.txt_display_e.insert(
            tk.END,
            f"N = {self.n}\nE = {self.e}"
        )
        self.txt_display_d.delete("1.0", tk.END)
        self.txt_display_d.insert(
            tk.END,
            f"N = {self.n}\nD = {self.d}"
        )
        self.txt_t3_n.delete(0, tk.END)
        self.txt_t3_n.insert(0, str(self.n))
        self.txt_t3_e.delete(0, tk.END)
        self.txt_t3_e.insert(0, str(self.e))
        messagebox.showinfo(
            "Thành công",
            "Tạo cặp khóa RSA thành công!"
        )

    def luu_cap_khoa(self):
        if self.n == 0:
            messagebox.showwarning(
                "Chưa có khóa",
                "Vui lòng tạo cặp khóa RSA trước!"
            )
            return
        path = filedialog.asksaveasfilename(
            title="Lưu cặp khóa RSA",
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt")]
        )
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write("===== CAP KHOA RSA =====\n")
                f.write(f"N={self.n}\n")
                f.write(f"E={self.e}\n")
                f.write(f"D={self.d}\n")
            messagebox.showinfo(
                "Thành công",
                "Đã lưu cặp khóa RSA thành công!"
            )
        except Exception as err:
            messagebox.showerror(
                "Lỗi",
                f"Không thể lưu file!\n{err}"
            )

    def mo_cap_khoa(self):
        path = filedialog.askopenfilename(
            title="Chọn file khóa RSA",
            filetypes=[("Text Files", "*.txt")]
        )
        if not path:
            return
        try:
            khoa = {}
            with open(path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line.startswith("N="):
                        khoa["N"] = int(line.split("=", 1)[1])
                    elif line.startswith("E="):
                        khoa["E"] = int(line.split("=", 1)[1])
                    elif line.startswith("D="):
                        khoa["D"] = int(line.split("=", 1)[1])
            if "N" not in khoa or "E" not in khoa or "D" not in khoa:
                raise ValueError
            self.n = khoa["N"]
            self.e = khoa["E"]
            self.d = khoa["D"]
            self.txt_e.delete(0, tk.END)
            self.txt_e.insert(0, str(self.e))
            self.txt_t3_n.delete(0, tk.END)
            self.txt_t3_n.insert(0, str(self.n))
            self.txt_t3_e.delete(0, tk.END)
            self.txt_t3_e.insert(0, str(self.e))
            self.txt_display_n.delete("1.0", tk.END)
            self.txt_display_n.insert(tk.END, str(self.n))
            self.txt_display_e.delete("1.0", tk.END)
            self.txt_display_e.insert(
                tk.END,
                f"N = {self.n}\nE = {self.e}"
            )
            self.txt_display_d.delete("1.0", tk.END)
            self.txt_display_d.insert(
                tk.END,
                f"N = {self.n}\nD = {self.d}"
            )
            messagebox.showinfo("Thành công", "Đã nạp cặp khóa từ file!")
        except Exception as err:
            messagebox.showerror(
                "Lỗi",
                f"File không đúng định dạng!\n\n{err}"
            )
    def tai_van_ban_can_ky(self):
        path = filedialog.askopenfilename(filetypes=[("All Files", "*.txt *.docx *.xlsx *.pdf"), ("Text Files", "*.txt")])
        if path:
            self.txt_source_text.delete("1.0", tk.END)
            self.txt_source_text.insert(tk.END, ThuatToanRSA.doc_noi_dung_file(path))

    def luu_van_ban_goc(self):
        t = self.txt_source_text.get("1.0", tk.END).strip()
        if not t: return
        file_types = [("Text File", "*.txt"), ("Word Document", "*.docx"), ("PDF Document", "*.pdf")]
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=file_types)
        if not path: return
        ext = path.split(".")[-1].lower()
        if ext == "txt":
            with open(path, "w", encoding="utf-8") as f: f.write(t)
        elif ext == "docx":
            from docx import Document
            doc = Document()
            doc.add_paragraph(t)
            doc.save(path)
        elif ext == "pdf":
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(path)
            c.drawString(100, 750, t) # Đơn giản hóa việc in text vào PDF
            c.save()
        messagebox.showinfo("Thành công", f"Đã lưu thành công định dạng .{ext}")

    def luu_chu_ky(self):
        t = self.txt_output_sig.get("1.0", tk.END).strip()
        if not t: return
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
        if path:
            with open(path, "w", encoding="utf-8") as f: f.write(t)

    def ky_so_van_ban(self):
        if self.n == 0 or self.d == 0:
            messagebox.showwarning("Lỗi khóa", "Vui lòng tạo hệ thống cặp khóa tại Tab 1 trước!")
            return
        chuoi_ky = self.txt_source_text.get("1.0", tk.END).strip()
        if not chuoi_ky:
            messagebox.showwarning("Thiếu dữ liệu", "Vui lòng nhập hoặc tải văn bản trước khi ký số!")
            self.txt_source_text.focus_set()
            return
        hash_bytes = hashlib.sha256(chuoi_ky.encode("utf-8")).digest()
        m = int.from_bytes(hash_bytes, byteorder='big')
        if m >= self.n:
            messagebox.showerror("Lỗi", "Khóa N hiện tại quá nhỏ so với hàm băm SHA-256 (256-bit). Hãy sinh khóa mới với P, Q lớn hơn!")
            return
        s_ket_qua = pow(m, self.d, self.n)
        self.original_text = chuoi_ky
        self.original_sig = s_ket_qua
        self.txt_output_sig.delete("1.0", tk.END)
        self.txt_output_sig.insert(tk.END, str(s_ket_qua))
        t_str = datetime.now().strftime("%H:%M:%S")
        self.dem_ls += 1
        self.data_ls[self.dem_ls] = {"text": chuoi_ky, "S": str(s_ket_qua)}
        f_row = ctk.CTkFrame(self.scroll_history, fg_color="transparent", height=32)
        f_row.pack(fill="x", pady=2)
        f_row.grid_propagate(False)
        f_row.columnconfigure(2, weight=1)
        btn_select = ctk.CTkButton(f_row, text="", fg_color="transparent", hover_color=("#D1D5DB", "#374151"), corner_radius=5, command=lambda idx=self.dem_ls: self.chon_lich_su_ky(idx))
        btn_select.place(x=0, y=0, relwidth=1, relheight=1)
        lbl_stt = ctk.CTkLabel(f_row, text=str(self.dem_ls), width=50)
        lbl_stt.grid(row=0, column=0, padx=5, sticky="w")
        lbl_time = ctk.CTkLabel(f_row, text=t_str, width=100)
        lbl_time.grid(row=0, column=1, padx=5, sticky="w")
        trich_doan = (chuoi_ky if len(chuoi_ky) < 40 else chuoi_ky[:37] + "...").replace("\n", " ")
        lbl_text = ctk.CTkLabel(f_row, text=trich_doan, anchor="w")
        lbl_text.grid(row=0, column=2, padx=5, sticky="ew")
        lbl_sig = ctk.CTkLabel(f_row, text=str(s_ket_qua)[:35] + "...", width=250, anchor="w")
        lbl_sig.grid(row=0, column=3, padx=5, sticky="w")
        for item in [lbl_stt, lbl_time, lbl_text, lbl_sig]:
            item.bind("<Button-1>", lambda event, idx=self.dem_ls: self.chon_lich_su_ky(idx))
        self.txt_t3_text.delete("1.0", tk.END); self.txt_t3_text.insert(tk.END, chuoi_ky)
        self.txt_t3_sig.delete("1.0", tk.END); self.txt_t3_sig.insert(tk.END, str(s_ket_qua))

    def chon_lich_su_ky(self, stt):
        if stt in self.data_ls:
            saved = self.data_ls[stt]
            self.txt_source_text.delete("1.0", tk.END); self.txt_source_text.insert(tk.END, saved["text"])
            self.txt_output_sig.delete("1.0", tk.END); self.txt_output_sig.insert(tk.END, saved["S"])
            self.txt_t3_text.delete("1.0", tk.END); self.txt_t3_text.insert(tk.END, saved["text"])
            self.txt_t3_sig.delete("1.0", tk.END); self.txt_t3_sig.insert(tk.END, saved["S"])

    def tai_van_ban_xac_minh(self):
        path = filedialog.askopenfilename(filetypes=[("All Files", "*.*")])
        if path:
            self.txt_t3_text.delete("1.0", tk.END); self.txt_t3_text.insert(tk.END, ThuatToanRSA.doc_noi_dung_file(path))
            
    def tai_chu_ky_xac_minh(self):
        path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
        if path:
            with open(path, "r", encoding="utf-8") as f: sig = f.read().strip()
            self.txt_t3_sig.delete("1.0", tk.END); self.txt_t3_sig.insert(tk.END, sig)
            
    def xac_minh_chu_ky(self):
        self.txt_log.delete("1.0", tk.END)
        self.txt_log.insert(tk.END, "⚙️ ĐANG THẨM ĐỊNH CHI TIẾT...\n\n")
        str_n = self.txt_t3_n.get().strip()
        str_e = self.txt_t3_e.get().strip()
        van_ban_ht = self.txt_t3_text.get("1.0", tk.END).strip()
        chu_ky_ht_raw = self.txt_t3_sig.get("1.0", tk.END).strip()

        if not str_n: messagebox.showerror("Lỗi", "Chưa nhập Khóa N!"); return
        if not str_e: messagebox.showerror("Lỗi", "Chưa nhập Chỉ số E!"); return
        if not van_ban_ht: messagebox.showerror("Lỗi", "Văn bản trống!"); return
        if not chu_ky_ht_raw: messagebox.showerror("Lỗi", "Chữ ký trống!"); return
        
        try:
            n = int(str_n)
            e = int(str_e)
            s_ht = int(''.join(filter(str.isdigit, chu_ky_ht_raw)))
        except:
            messagebox.showerror("Lỗi", "Dữ liệu N, E, S phải là số nguyên!"); return

        goc_van_ban = getattr(self, 'original_text', van_ban_ht)
        goc_chu_ky = getattr(self, 'original_sig', s_ht)

        van_ban_bi_sua = (van_ban_ht != goc_van_ban)
        chu_ky_bi_sua = (s_ht != goc_chu_ky)

        hash_bytes = hashlib.sha256(van_ban_ht.encode("utf-8")).digest()
        hash_ht = int.from_bytes(hash_bytes, byteorder='big')
        m2 = pow(s_ht, e, n)
        toan_hoc_dung = (hash_ht == m2)
        if van_ban_bi_sua and chu_ky_bi_sua:
            messagebox.showerror("Cảnh báo", "Văn bản không toàn vẹn và chữ ký không hợp lệ!")
            self.txt_log.insert(tk.END, "❌ Kết quả: Văn bản & Chữ ký đều đã bị sửa.")

        elif van_ban_bi_sua:
            messagebox.showwarning("Cảnh báo", "Văn bản không toàn vẹn!")
            self.txt_log.insert(tk.END, "⚠️ Kết quả: Văn bản đã bị sửa.")

        elif chu_ky_bi_sua:
            messagebox.showerror("Lỗi", "Chữ ký không hợp lệ!")
            self.txt_log.insert(tk.END, "❌ Kết quả: Chữ ký đã bị sửa.")

        elif not toan_hoc_dung:
            messagebox.showerror("Lỗi", "Chữ ký không khớp với văn bản!")
            self.txt_log.insert(tk.END, "❌ Kết quả: Dữ liệu không khớp thuật toán RSA.")

        else:
            messagebox.showinfo("Thành công", "Chữ ký hợp lệ, văn bản toàn vẹn!")
            self.txt_log.insert(tk.END, "✅ Kết quả: Hợp lệ.")

if __name__ == "__main__":
    root = ctk.CTk()
    app = UngDungChuKySoRSA(root)
    root.mainloop()